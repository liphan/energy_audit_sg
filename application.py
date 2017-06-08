
# coding: utf-8

# In[4]:

import os

import sys
if sys.version_info[0] < 3:        # check python version
    from StringIO import StringIO  # https://stackoverflow.com/questions/22604564/how-to-create-a-pandas-dataframe-from-string
else:
    from io import StringIO

import pandas as pd

import matplotlib
matplotlib.use('Agg')
from matplotlib import pyplot as plt
from matplotlib import dates
import matplotlib.dates as mdates
import seaborn as sns

from datetime import datetime
from datetime import timedelta

from docx import Document
from docx.shared import Inches
document = Document()
run = document.add_paragraph().add_run()
font = run.font


# In[5]:

# %load user.py


# In[66]:
# import psycopg2

from flask import Flask, render_template, flash, request, url_for, redirect, session, send_file
from flask_wtf import Form

# from flask_wtf import FlaskForm
# from flask_wtf.file import FileField, FileRequired, FileAllowed

from wtforms import TextField, TextAreaField, SubmitField, validators, ValidationError, PasswordField, StringField
from wtforms import BooleanField
#from wtforms_alchemy import PhoneNumberField

#from forms import ContactForm, SignupForm
from flask_mail import Message, Mail
#from passlib.hash import sha256_crypt
#import escape-string as thwart
#import gc


# In[67]:

from flask import Flask

app = Flask(__name__)

app.secret_key = 'development key'

app.config["MAIL_SERVER"] = "smtp.gmail.com"
app.config["MAIL_PORT"] = 465
app.config["MAIL_USE_SSL"] = True
app.config["MAIL_USERNAME"] = 'mgresap@gmail.com'      # 'contact@example.com'
app.config["MAIL_PASSWORD"] = 'wevxbgm1'               # 'your-password'


#from routes import mail
mail = Mail()

mail.init_app(app)

from flask_sqlalchemy import SQLAlchemy
from werkzeug import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

# app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://postgres:admin@localhost:5432/audit'
app.config['SQLALCHEMY_DATABASE_URI'] = 'postgres://ynmfoitwdrfufl:fcdc5a2eff67d02812cab4bda26e1bff1d4b050aa09b54c0b7c730ea93a5c0b5@ec2-54-221-255-153.compute-1.amazonaws.com:5432/d6rmmhfa7ijsat?sslmode=require'   # ?sslmode=require
#from models import db
db = SQLAlchemy(app)

# db.create_all()
# db.session.commit()
UPLOAD_FOLDER = '/path/to/the/uploads'
ALLOWED_EXTENSIONS = set(['csv'])
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

class User(db.Model):
    __tablename__ = 'users'
#     uid = db.Column(db.Integer)
    firstname = db.Column(db.String(100))
    lastname = db.Column(db.String(100))
    email = db.Column(db.String(120), unique=True)
    mobile = db.Column(db.Integer, primary_key = True)
    pwdhash = db.Column(db.String(150))

    def __init__(self, firstname, lastname, email, mobile, password):

        self.firstname = firstname.title()
        self.lastname = lastname.title()
        self.email = email.lower()
        self.mobile = mobile
        self.set_password(password)

    def set_password(self, password):
        self.pwdhash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.pwdhash, password)

# db.init_app(app)

#import app.routes


# In[68]:

class SignupForm(Form):

    firstname = StringField("First name",  [validators.DataRequired("Please enter your first name.")])
    lastname = StringField("Last name",  [validators.DataRequired("Please enter your last name.")])
    email = StringField("Email",  [validators.DataRequired("Please enter your email address."),
                                 validators.Email("Please enter your email address.")])
    #password = PasswordField('Password', [validators.Required("Please enter a password.")])
    mobile = StringField('Mobile', [validators.DataRequired(), validators.Length(min=8, max=8)])#country_code='SG', display_format='national')

    password = PasswordField('New Password', [validators.DataRequired("Please enter a password.")])#,
#                                               validators.EqualTo('confirm', message='Passwords must match')])
#     confirm = PasswordField('Repeat Password')
#     accept_tos = BooleanField('I accept the Terms of Service and Privacy Notice (updated Jan 22, 2015)',
#                               [validators.DataRequired()])

    submit = SubmitField("Create account")


# class RegistrationForm(Form):
#     firstname = TextField("First name",  [validators.Required("Please enter your first name.")])
#     lastname = TextField("Last name",  [validators.Required("Please enter your last name.")])
#     email = TextField("Email address",  [validators.Required("Please enter your email address."),
#                                          validators.Email("Please enter your email address.")])
#     #username = TextField('Username', [validators.Length(min=4, max=20)])
#     #email = TextField('Email Address', [validators.Length(min=6, max=50)])
#     password = PasswordField('New Password', [validators.Required("Please enter a password."),
#                                               validators.EqualTo('confirm', message='Passwords must match')])
#     confirm = PasswordField('Repeat Password')
#     accept_tos = BooleanField('I accept the Terms of Service and Privacy Notice (updated Jan 22, 2015)',
#                               [validators.Required()])
    #submit = SubmitField("Create account")

    def __init__(self, *args, **kwargs):
        Form.__init__(self, *args, **kwargs)

    def validate(self):
        if not Form.validate(self):
            return False

        user = User.query.filter_by(email = self.email.data.lower()).first()
        if user:
            self.email.errors.append("That email is already taken")
            return False
        else:
            return True


class SigninForm(Form):
    email = TextField("Email",  [validators.DataRequired("Please enter your email address."), validators.Email("Please enter your email address.")])
    password = PasswordField('Password', [validators.DataRequired("Please enter a password.")])
    submit = SubmitField("Sign In")

    def __init__(self, *args, **kwargs):
        Form.__init__(self, *args, **kwargs)

    def validate(self):
        if not Form.validate(self):
            return False

        user = User.query.filter_by(email = self.email.data.lower()).first()
        if user and user.check_password(self.password.data):
            return True
        else:
            self.email.errors.append("Invalid e-mail or password")
            return False


class ContactForm(Form):
    name = TextField("Name",  [validators.DataRequired("Please enter your name.")])
    email = TextField("Email",  [validators.DataRequired("Please enter your email address."), validators.Email("Please enter your email address.")])
    subject = TextField("Subject",  [validators.DataRequired("Please enter a subject.")])
    message = TextAreaField("Message",  [validators.DataRequired("Please enter a message.")])
    submit = SubmitField("Send")


# class UploadForm(FlaskForm):
#     datafile = FileField('your csv', validators=[FileRequired(), FileAllowed(['csv'], 'csv file only!')])
#
#     def __init__(self, *args, **kwargs):
#         Form.__init__(self, *args, **kwargs)
#
#     def validate(self):
#         if not Form.validate(self):
#             return False

# In[69]:

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    form = ContactForm()

    if request.method == 'POST':
        if form.validate() == False:
            flash('All fields are required.')
            return render_template('contact.html', form=form)
        else:
            msg = Message(form.subject.data, sender='mgresap@gmail.com', recipients=['lip_h@hotmail.com'])
            msg.body = """
            From: %s &lt;%s&gt;
            %s
            """ % (form.name.data, form.email.data, form.message.data)

            mail.send(msg)

            return 'Form posted.'

    elif request.method == 'GET':
        return render_template('contact.html', form=form)


# In[70]:

# from flask import Flask

# app = Flask(__name__)

#mail = Mail()
@app.route('/signup', methods=['GET', 'POST'])
def signup():

    form = SignupForm()

    if 'email' in session:
        return redirect(url_for('entries'))
    #     if 'email' not in session:
    #         return redirect(url_for('signin'))
    #
    #     user = User.query.filter_by(email = session['email']).first()
    #
    #     if user is None:
    #         return redirect(url_for('signin'))
    #     else:
    # #         return render_template('profile.html')
    #         return redirect(url_for('upload'))


    if request.method == 'POST':

        if form.validate() == False:

            return render_template('signup.html', form=form)
        else:

            newuser = User(form.firstname.data, form.lastname.data, form.email.data, form.mobile.data,
                           form.password.data)
            db.session.add(newuser)
            db.session.commit()
            session['email'] = newuser.email
#             return "[1] Create a new user [2] sign in the user [3] redirect to the user's profile"
            return redirect(url_for('entries'))
        #     if 'email' not in session:
        #         return redirect(url_for('signin'))
        #
        #     user = User.query.filter_by(email = session['email']).first()
        #
        #     if user is None:
        #         return redirect(url_for('signin'))
        #     else:
        # #         return render_template('profile.html')
        #         return redirect(url_for('upload'))

    else:
        return render_template('signup.html', form=form)


# In[71]:

#  url mapping for /profile



# @app.route('/profile', methods=['GET', 'POST'])
# def profile():
#     if 'email' not in session:
#         return redirect(url_for('signin'))
#
#     user = User.query.filter_by(email = session['email']).first()
#
#     if user is None:
#         return redirect(url_for('signin'))
#     else:
# #         return render_template('profile.html')
#         return redirect(url_for('upload'))
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1] in ALLOWED_EXTENSIONS

# @app.route('/upload', methods=['GET', 'POST'])
# def upload():
#     if 'email' not in session:
#         return redirect(url_for('signin'))
#
#     user = User.query.filter_by(email = session['email']).first()
#
#     if user is None:
#         return redirect(url_for('signin'))
#     else:
#         if request.method == 'POST':
#             f = request.files['file_location']
#             if f and allowed_file(f.filename):
#                 filename = secure_filename(f.filename)
#                 # f.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
#             return redirect(url_for('entries'))


# @app.route('/upload', methods=['GET', 'POST'])
# def upload():
#     if 'email' not in session:
#         return redirect(url_for('signin'))
#
#     user = User.query.filter_by(email = session['email']).first()
#
#     if user is None:
#         return redirect(url_for('signin'))
#     else:
# #         return render_template('profile.html')
#         # return redirect(url_for('upload'))
#
#         if request.method == 'POST':
#             f = request.files['file_location']
#             # path_ = f.read()
#             return redirect(url_for('entries'))

        # if form.validate_on_submit():
        #     f = form.datafile.data
        #     filename = secure_filename(f.filename)
        #     f.save(os.path.join(
        #         app.instance_path, 'photos', filename
        #     ))
        #     return redirect(url_for('entries'))
        #
        # return render_template('upload.html', form=form)


@app.route('/entries', methods=['GET', 'POST'])
def entries():
    if 'email' not in session:
        return redirect(url_for('signin'))

    user = User.query.filter_by(email = session['email']).first()

    if user is None:
        return redirect(url_for('signin'))
    else:
#         return render_template('profile.html')
        # return redirect(url_for('entries'))

        if request.method == "POST":
            f = request.files['file_location']
            if f and allowed_file(f.filename):
                filename = secure_filename(f.filename)
                # file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                file_content = f.read()
            # path = request.form['file location']
            # path = request.form['file location']
            ref = request.form['project ref']
            building_name = request.form['building name']
            building_address = request.form['building address']
            post_code = request.form['postal code']
            building_type = request.form['building type']
            building_age = request.form['building age']
            last_submit = request.form['last audit date']
            gfa = request.form['gross floor area']
            ac_area = request.form['air-conditioned area']
            room_qty = request.form['guest rooms']
            auditor = request.form['auditor']
            owner = request.form['owner']

            plant_loc = request.form['plant loc']
            audit_begin = request.form['audit start']
            audit_finish = request.form['audit end']
            notice = request.form['notice']
            submission = request.form['submission']

            chiller_cap = request.form['chiller cap']
            chiller_operating_qty = int(request.form['chiller qty'])
            chwp_rated_flow = float(request.form['chwp flow'])
            chwp_rated_head = float(request.form['chwp head'])
            chwp_operating_qty = int(request.form['chwp qty'])
            chwp_motor_effy = float(request.form['chwp motor effy'])
            chwp_pump_effy = float(request.form['chwp pump effy'])
            cwp_rated_flow = float(request.form['cwp flow'])
            cwp_rated_head = float(request.form['cwp head'])
            cwp_operating_qty = int(request.form['cwp qty'])
            cwp_motor_effy = float(request.form['cwp motor effy'])
            cwp_pump_effy = float(request.form['cwp pump effy'])
            ct_cap = float(request.form['ct cap'])
            ct_operating_qty = int(request.form['ct qty'])
            fan_qty = int(request.form['ct fan qty'])
            flow_per_fan = float(request.form['flow each fan'])
            absorbed_fan_power = float(request.form['absorbed fan power'])

            full_fan_power = absorbed_fan_power * fan_qty * ct_operating_qty
            total_ct_cap = ct_cap * ct_operating_qty




    # prepare pandas dataframe for plots
            # directory = os.path.dirname(path)
            # path_ = os.path.join(path)
            # abs_path = os.path.abspath(path_)
            # with open(abs_path) as f:
            DATA = StringIO(file_content)
            # df = pd.read_csv(file_content)
            df = pd.read_csv(DATA, sep=',', header=0, parse_dates=True, tupleize_cols=False, error_bad_lines=False, warn_bad_lines=True, skip_blank_lines=True)

            # df = pd.read_csv(abs_path)
            directory = os.path.dirname('file_location')

            df['date'] = pd.to_datetime(df['date'])
            df['time'] = pd.to_datetime(df['time']).dt.strftime('%H:%M:%S')
            audit_start = df['date'][0]
            audit_end = df['date'].iloc[-1]
            df['day of week'] = df['date'].dt.weekday_name
            df['heat gain kw'] = df['chw l/sec']*4.19*(df['chwr temp']-df['chws temp'])
            df['heat rejected kw'] = df['cw l/sec']*4.19*(df['cwr temp']-df['cws temp'])

            df['heat gain rt'] = 0.284333*df['chw l/sec']*4.19*(df['chwr temp']-df['chws temp'])
            df['heat rejected rt'] = 0.284333*df['cw l/sec']*4.19*(df['cwr temp']-df['cws temp'])

            df['% heat balance'] = ((df['chw l/sec']*4.19*(df['chwr temp']-df['chws temp'])+
                                     df['chiller kwe']) - df['cw l/sec']*4.19*(df['cwr temp']-
                                    df['cws temp']))*100/(df['cw l/sec']*4.19*(df['cwr temp']-df['cws temp']))

            total_kwe = df['chiller kwe'].sum()
            total_cooling = round(df['heat gain rt'].sum(), 2)
            total_heat_reject = round(df['heat rejected rt'].sum(), 2)
            plant_effy = round((total_kwe / total_cooling), 2)
            total_count = len(df)
            results = df['% heat balance']
            success_count = 0
            for result in results:
                if result > -5 and result < 5:
                    success_count += 1
            exceed_5_percent_count = 0
            for result in results:
                if result > 5:
                    exceed_5_percent_count += 1
            below_minus_5_percent_count = 0
            for result in results:
                if result < -5:
                    below_minus_5_percent_count += 1
            percent_heatbal_plusminus5percent = round((float(success_count) / total_count) * 100)

            days = df['day of week'].unique()
            dfp = df.pivot(index='time', columns='day of week', values='heat gain rt')
            day_ave = 0
            for day in days:
                day_ave = day_ave + dfp[day].mean()
            ave = day_ave / len(days)


            df['chiller kw/rt'] = df['chiller kwe'] / df['heat gain rt']

            df['chwp head'] = (((df['chw l/sec']/chwp_operating_qty)/chwp_rated_flow)**2) * chwp_rated_head


            df['chwp kwe'] = (df['chw l/sec']*9810*df['chwp head']/chwp_operating_qty)/((chwp_motor_effy/100) *
                                                                                        (chwp_pump_effy/100) * 10**6)
            df['chwp kw/rt'] = df['chwp kwe'] / df['heat gain rt']

            df['cwp head'] = (((df['cw l/sec']/cwp_operating_qty)/cwp_rated_flow)**2) * cwp_rated_head
            df['cwp kwe'] = (df['cw l/sec']*9810*df['cwp head']/cwp_operating_qty)/((cwp_motor_effy/100) *
                                                                                    (cwp_pump_effy/100) * 10**6)
            df['cwp kw/rt'] = df['cwp kwe'] / df['heat gain rt']

            df['ct_cap_ratio_needed'] = df['heat rejected kw'] / total_ct_cap
            df['fan power needed'] = full_fan_power * df['ct_cap_ratio_needed']
            df['ct kw/rt'] = df['fan power needed'] / df['heat gain rt']

            df['plant effy'] = (df['chiller kwe'] + df['chwp kwe']*chwp_operating_qty
                                      + df['cwp kwe']*cwp_operating_qty + df['fan power needed']) / df['heat gain rt']

    # cooling load profile plot

    #         %matplotlib inline

            sns.set_context("notebook", font_scale=1.2, rc={"lines.linewidth": 1.2})
            sns.set_style ("whitegrid") #("darkgrid")
            custom_style = {
                        'grid.color': '0.8',
                        'grid.linestyle': '--',
                        'grid.linewidth': 0.5,
            }
            sns.set_style(custom_style)



            fig, ax = plt.subplots(1, 1,
                                   figsize=(9.5, 6))
            xfmt = mdates.DateFormatter('%H:%M:%S')
            ax.xaxis.set_major_formatter(xfmt)
            ax.xaxis_date()

            ax.plot(dfp)

            ax.set(xlabel='Time',
                   ylabel='Cooling Load (RT)',
                   title=('Cooling Load Profile [%s - %s]')%(audit_start, audit_end))

            ax.legend(dfp.columns, loc='best') #loc=2
            fig.autofmt_xdate()

            # Draw a horizontal line showing the mean
            plt.axhline(y=ave, xmin=0.0, xmax=1.00, color='r', linestyle='dashed',
                            label='ave', lw=1.5)

            # draw VERTICAL lines
            plt.axvline(x='15:07:00',  ymin=0.0, ymax = 0.9, linewidth=2, color='r', linestyle='dashed')



            ax.annotate(('ave(10am-9pm) = %.2f'+' '+'RTon')%ave, xy=('15:07:00', 586), xytext=('15:00:15', ave+2),

                        )
            path_clp = os.path.join(directory, 'cooling load profile.png')
            plt.savefig(path_clp)

    #         chw temp plot

            dfp_chws = df.pivot(index='time', columns='day of week', values='chws temp')
            dfp_chws = dfp_chws.rename(columns = {'Monday':'CHWS Mon','Saturday':'CHWS Sat','Sunday':'CHWS Sun','Tuesday':'CHWS Tues',
                                                 'Wednesday':'CHWS Wed','Thursday':'CHWS Thurs','Friday':'CHWS Fri'})
            dfp_chwr = df.pivot(index='time', columns='day of week', values='chwr temp')
            dfp_chwr = dfp_chwr.rename(columns = {'Monday':'CHWR Mon','Saturday':'CHWR Sat','Sunday':'CHWR Sun','Tuesday':'CHWR Tues',
                                                 'Wednesday':'CHWR Wed','Thursday':'CHWR Thurs','Friday':'CHWR Fri'})
            dfp_chw = pd.concat([dfp_chws, dfp_chwr], axis=1)
            dfp_chw

            fig, ax = plt.subplots(1, 1, figsize=(9.5, 6))
            xfmt = mdates.DateFormatter('%H:%M:%S')
            ax.xaxis.set_major_formatter(xfmt)
            ax.xaxis_date()

            ax.plot(dfp_chw)
            ax.set_xlabel('Time')
            ax.set_ylabel('Temperature (degC)')
            ax.set_title('CHW Temperature')
            ax.legend(dfp_chw.columns, loc='best')

            path_chwt = os.path.join(directory, 'chw temp.png')
            plt.savefig(path_chwt)


    # report starts here !

            document.add_heading('ENERGY AUDIT REPORT FOR BUILIDNG COOLING SYSTEM', level=1).alignment=1
            document.add_heading('FOR', level=1).alignment=1
            document.add_heading('%s'%building_name, level=1).alignment=1
            document.add_heading('%s'%building_address, level=1).alignment=1

            document.add_heading('', level=1)
            document.add_heading('', level=1)
            document.add_heading('', level=1)
            document.add_heading('', level=1)
            document.add_heading('', level=1)
            document.add_heading('', level=1)

            document.add_paragraph('Submitted By').alignment=1
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('Registration No.').alignment=1

            document.add_paragraph('')
            document.add_paragraph('')
    # new page
            document.add_heading('Contents', level=1)
            document.add_paragraph('1.0  Executive Summary & Recommendation')
            document.add_paragraph('2.0  Building Information')
            document.add_paragraph('3.0  Energy Audit Information For Building Cooling System')
            document.add_paragraph('        3.1  Chilled Water Plant Design information')
            document.add_paragraph('        Table 1: Chiller Information')
            document.add_paragraph('        Table 2: Ancillary equipment Information')


            document.add_paragraph('        3.2  Chilled Water Plant Normal Operating Hours')
            document.add_paragraph('        3.3  Description of Plant Control Strategy')
            document.add_paragraph('4.0  Instrumentations')
            document.add_paragraph('        Table 3: Instrumentation Table')

            document.add_paragraph('5.0  Chiller Plant Performance Analysis (1 week data)')
            document.add_paragraph('        Fig 5.1  Super-imposed plot of 24 hr Cooling Load Profile RT')
            document.add_paragraph('        Fig 5.2  Histogram of Cooling Load Occurrences')
            document.add_paragraph('''        Fig 5.3  Super-imposed plot of daily chilled water supply/return
                           temperature degC''')
            document.add_paragraph('''        Fig 5.4  Super-imposed plot of daily chilled water temperature
                           difference degC''')

            document.add_paragraph('''        Fig 5.5  Super-imposed plot of daily condenser water supply/return
                           temperature degC''')
            document.add_paragraph('''        Fig 5.6  Super-imposed plot of daily condenser water temperature
                           difference degC''')
            document.add_paragraph('        Fig 5.7  Super-imposed plot of daily chilled water GPM/RT ')
            document.add_paragraph('        Fig 5.8  Super-imposed plot of daily condenser water GPM/RT')
            document.add_paragraph('        Fig 5.9  Cooling Tower Approach Temperature')
            document.add_paragraph('        Fig 5.10 Super-imposed plot of daily chiller efficiency kW/RT')
            document.add_paragraph('        Fig 5.11 Super-imposed plot of daily chilled water pump efficiency kW/RT')
            document.add_paragraph('        Fig 5.12 Super-imposed plot of daily condenser water pump efficiency kW/RT')
            document.add_paragraph('        Fig 5.13 Super-imposed plot of daily cooling tower efficiency kW/RT')
            document.add_paragraph('        Fig 5.14 Super-imposed plot of daily chiller plant system efficiency kW/RT')
            document.add_paragraph('        Fig 5.15 Scatter plot of chiller plant efficiency over cooling load')
            document.add_paragraph('        Fig 5.16 Scatter plot of chilled water pump efficiency over cooling load')
            document.add_paragraph('        Fig 5.17 Scatter plot of condenser water pump efficiency over cooling load')
            document.add_paragraph('        Fig 5.18 Scatter plot of cooling tower efficiency over cooling load')
            document.add_paragraph('        5.1      Summary of Chilled Water Plant Operating Performance')
            document.add_paragraph('        Table 4: Chilled Water Plant Performance Summary')

            document.add_paragraph('6.0  Summary of Heat Balance')
            document.add_paragraph('        Fig 6.1  System Level Heat Balance Plot')
            document.add_paragraph('        Table 5: Heat Balance Summary')
            document.add_paragraph('7.0  Schedule of space operating conditions')
            document.add_paragraph('        Table 6: Space Condition Schedule')
            document.add_paragraph('APPENDIX')
            document.add_paragraph('Checklist of Plant Operating Condition (for best practices)')
            document.add_paragraph('     Table 7: Checklist of Plant Operating Condition')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')



    #  new page
            document.add_heading('1.0     Executive Summary & Recommendation', level=1)
            document.add_paragraph('')
            document.add_paragraph(('This report highlights the findings and recommendations obtained from the energy audit     performed at %s from %s to %s for 24 hrs.') %(building_name, audit_start, audit_end))
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('Recommendations for maintenance improvements and low cost energy conservation measures.')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('Recommendations which would incur capital expenditure.')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
            document.add_paragraph('')
    # new page
            document.add_heading('2.0     Building Information', level=1)
            document.add_paragraph('')

            bldg_info = document.add_table(rows=10, cols=2)

            bldg_info.rows[0].cells[0].text = 'Project Reference Number'
            bldg_info.rows[0].cells[1].text = '%s' %ref
            bldg_info.rows[1].cells[0].text = 'Building Name'
            bldg_info.rows[1].cells[1].text = '%s' %building_name
            bldg_info.rows[2].cells[0].text = 'Building Address'
            bldg_info.rows[2].cells[1].text = '%s' %building_address
            bldg_info.rows[3].cells[0].text = 'Postal Code'
            bldg_info.rows[3].cells[1].text = '%s' %post_code
            bldg_info.rows[4].cells[0].text = 'Building Type'
            bldg_info.rows[4].cells[1].text = '%s' %building_type
            bldg_info.rows[1].cells[0].text = 'Building Age'
            bldg_info.rows[1].cells[1].text = '%s' %building_age
            bldg_info.rows[2].cells[0].text = 'Date of last Energy Audit Submission'
            bldg_info.rows[2].cells[1].text = '%s' %last_submit
            bldg_info.rows[3].cells[0].text = 'Gross floor area (GFA), m2'
            bldg_info.rows[3].cells[1].text = '%s' %gfa
            bldg_info.rows[4].cells[0].text = 'Air conditioned area, m2'
            bldg_info.rows[4].cells[1].text = '%s' %ac_area
            bldg_info.rows[5].cells[0].text = 'Number of guest rooms (for hotels/service apartments)'
            bldg_info.rows[5].cells[1].text = '%s' %room_qty
    # new page
            document.add_heading('3.0     Energy Audit Information For Building Cooling System', level=1)
            document.add_paragraph('')
            document.add_paragraph(('''%s was appointed by %s, owner of %s to be the Energy Auditor for the 3 yearly     submission of the operating system efficiency (OSE) of the centralized Chilled Water Plant. The report will     present the performance of centralized Chilled Water Plant efficiency based on the measurements from the     permanent instrumentations installed on site.''')%(auditor,owner,building_name))

            document.add_paragraph('')

            audit_info = document.add_table(rows=7, cols=2)

            audit_info.rows[0].cells[0].text = 'Location'
            audit_info.rows[0].cells[1].text = '%s' %plant_loc

            audit_info.rows[1].cells[0].text = 'Energy Audit Period'
            audit_info.rows[1].cells[1].text = '%s to %s' %(audit_begin,audit_finish)

            audit_info.rows[2].cells[0].text = 'Date of notice served'
            audit_info.rows[2].cells[1].text = '%s' %notice
            audit_info.rows[3].cells[0].text = 'Date of submission in notice'
            audit_info.rows[3].cells[1].text = '%s' %submission
            audit_info.rows[4].cells[0].text = 'Data Logging Interval '
            audit_info.rows[4].cells[1].text = '1 minute sampling'
            audit_info.rows[5].cells[0].text = 'Trend Logged Parameters'
            audit_info.rows[5].cells[1].text = '                               '
    # new page
            document.add_heading('5.0      Chiller Plant Performance Analysis (1 week data)', level=1)
            document.add_paragraph('')
            document.add_picture(path_clp, width=Inches(6.00))
            os.remove(path_clp)

            document.add_paragraph('')
            document.add_picture(path_chwt, width=Inches(6.00))
            os.remove(path_chwt)

    # new page

            document.add_heading('6.0      Summary of Heat Balance', level=1)
    #             document.add_picture('heat balance percent.png', width=Inches(6.00))
    #             document.add_paragraph('System Level Heat Balance Plot').alignment=1  # 0 for left, 1 for center, 2 for right


            table = document.add_table(rows=10, cols=3)

            row0 = table.rows[0]
            row0.cells[1].text = 'Quantity'
            row0.cells[2].text = 'Unit'

            row1 = table.rows[1]
            row1.cells[0].text = 'Sum of total electrical energy used'
            row1.cells[1].text = '%d' %total_kwe
            row1.cells[2].text = 'kWh'

            row2 = table.rows[2]
            row2.cells[0].text = 'Sum of total cooling produced'
            row2.cells[1].text = '%d' %total_cooling
            row2.cells[2].text = 'RTh'

            row3 = table.rows[3]
            row3.cells[0].text = 'Sum of total heat rejected'
            row3.cells[1].text = '%d' %total_heat_reject
            row3.cells[2].text = 'RTh'

            row4 = table.rows[4]
            row4.cells[0].text = 'Chiller Plant Efficiency'
            row4.cells[1].text = '%s' %total_cooling
            row4.cells[2].text = 'kW/RT'

            row5 = table.rows[5]
            row5.cells[0].text = 'Total Heat Balance Data Count'
            row5.cells[1].text = '%d' %total_count
            row5.cells[2].text = '-'

            row6 = table.rows[6]
            row6.cells[0].text = 'Data Count > + 5% error '
            row6.cells[1].text = '%d' %exceed_5_percent_count
            row6.cells[2].text = '-'

            row7 = table.rows[7]
            row7.cells[0].text = 'Data Count < - 5% error '
            row7.cells[1].text = '%d' %below_minus_5_percent_count
            row7.cells[2].text = '-'

            row8 = table.rows[8]
            row8.cells[0].text = 'Data Count within +/-5% error '
            row8.cells[1].text = '%d' %success_count
            row8.cells[2].text = '-'

            row9 = table.rows[9]
            row9.cells[0].text = '% Heat Balance within +/-5% error'
            row9.cells[1].text = '%s' %percent_heatbal_plusminus5percent
            row9.cells[2].text = '%'

    # new page

            document.add_heading('7.0      Schedule of space operating conditions', level=1)

            # document.save(os.path.join(directory, 'report.docx'))
    #         document.save('report.odt')
            f = StringIO()      # https://stackoverflow.com/questions/27029933/generating-word-docs-with-flask
            document.save(f)    # http://flask.pocoo.org/snippets/32/
            length = f.tell()
            f.seek(0)
            return send_file(f, as_attachment=True, attachment_filename='report.docx')
        return render_template('entries.html')



# In[72]:

# url mapping for /signin

@app.route('/signin', methods=['GET', 'POST'])
def signin():
    form = SigninForm()

    if 'email' in session:
        return redirect(url_for('entries'))
    #     if 'email' not in session:
    #         return redirect(url_for('signin'))
    #
    #     user = User.query.filter_by(email = session['email']).first()
    #
    #     if user is None:
    #         return redirect(url_for('signin'))
    #     else:
    # #         return render_template('profile.html')
    #         return redirect(url_for('upload'))


    if request.method == 'POST':
        if form.validate() == False:
            return render_template('signin.html', form=form)
        else:
            session['email'] = form.email.data
            return redirect(url_for('entries'))
        #     if 'email' not in session:
        #         return redirect(url_for('signin'))
        #
        #     user = User.query.filter_by(email = session['email']).first()
        #
        #     if user is None:
        #         return redirect(url_for('signin'))
        #     else:
        # #         return render_template('profile.html')
        #         return redirect(url_for('entries'))

    elif request.method == 'GET':
        return render_template('signin.html', form=form)


# In[73]:

@app.route('/signout')
def signout():

  if 'email' not in session:
    return redirect(url_for('signin'))

  session.pop('email', None)
  return redirect(url_for('home'))


# In[ ]:




# In[27]:

# @app.route('/register/', methods=["GET","POST"])
# def register_page():
#     try:
#         form = RegistrationForm(request.form)

#         if request.method == "POST" and form.validate():
#             username  = form.username.data
#             email = form.email.data
#             password = sha256_crypt.encrypt((str(form.password.data)))
#             c, conn = connection()

#             x = c.execute("SELECT * FROM users WHERE username = (%s)",
#                           (thwart(username)))

#             if int(x) > 0:
#                 flash("That username is already taken, please choose another")
#                 return render_template('register.html', form=form)

#             else:
#                 c.execute("INSERT INTO users (username, password, email, tracking) VALUES (%s, %s, %s, %s)",
#                           (thwart(username), thwart(password), thwart(email), thwart("/introduction-to-python-programming/")))

#                 conn.commit()
#                 flash("Thanks for registering!")
#                 c.close()
#                 conn.close()
#                 gc.collect()

#                 session['logged_in'] = True
#                 session['username'] = username

#                 return redirect(url_for('dashboard'))

#         return render_template("register.html", form=form)

#     except Exception as e:
#         return(str(e))


# In[74]:



# In[ ]:


# In[6]:

if __name__ == "__main__":
#     app.debug=True
    app.run()
    # app.run(port=33507)


# In[ ]:
